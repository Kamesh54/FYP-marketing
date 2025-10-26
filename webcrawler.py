import requests
import json
import os
import uuid
import time
from datetime import datetime
from urllib.parse import urlparse, urljoin
from collections import deque
from bs4 import BeautifulSoup
import logging
from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse  # Fixed import
from pydantic import BaseModel, HttpUrl
from typing import Optional, List, Dict, Set
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(title="Web Crawler API", description="A powerful web crawler for extracting content")

# Request and response models
class CrawlRequest(BaseModel):
    start_url: HttpUrl
    delay: Optional[float] = 1.0
    timeout: Optional[int] = 10
    max_pages: Optional[int] = 10  # Added max_pages limit
    output_filename: Optional[str] = None

class CrawlResponse(BaseModel):
    job_id: str
    status: str
    message: str
    pages_crawled: Optional[int] = None
    word_document: Optional[str] = None
    json_file: Optional[str] = None

# In-memory job storage (use Redis/database in production)
job_status = {}

class WebCrawler:
    def __init__(self, delay: float = 1.0, timeout: int = 10, max_pages: int = 10):
        self.delay = delay
        self.timeout = timeout
        self.max_pages = max_pages
        self.visited_urls: Set[str] = set()
        self.pending_urls: List[str] = []
        self.extracted_content: List[Dict] = []
        self.base_domain: Optional[str] = None
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36"
        }

    def _is_same_domain(self, url: str) -> bool:
        """Check if URL belongs to the same domain."""
        try:
            parsed_url = urlparse(url)
            return parsed_url.netloc == self.base_domain
        except:
            return False

    def _clean_url(self, url: str, base_url: str) -> str:
        """Clean and normalize URL."""
        try:
            absolute_url = urljoin(base_url, url)
            parsed = urlparse(absolute_url)
            clean_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
            if parsed.query:
                clean_url += f"?{parsed.query}"
            return clean_url
        except:
            return ""

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type(requests.RequestException)
    )
    def _extract_page_content(self, url: str) -> Dict:
        """Extract content from a single page using requests and BeautifulSoup."""
        logger.info(f"Extracting content from: {url}")
        try:
            response = requests.get(url, headers=self.headers, timeout=self.timeout)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            content = {
                'url': url,
                'title': soup.title.get_text(strip=True) if soup.title else '',
                'headers': self._extract_headers(soup),
                'paragraphs': self._extract_paragraphs(soup),
                'tables': self._extract_tables(soup),
                'links': self._extract_links(soup, url),
                'content': ''
            }

            content['content'] = " ".join(
                [h['text'] for h in content['headers']] +
                content['paragraphs']
            )

            return content
        except requests.RequestException as e:
            logger.error(f"Failed to extract content from {url}: {e}")
            return {
                'url': url,
                'title': '',
                'headers': [],
                'paragraphs': [],
                'tables': [],
                'links': [],
                'content': ''
            }

    def _extract_headers(self, soup: BeautifulSoup) -> List[Dict]:
        """Extract headers (h1-h6) from the page."""
        headers = []
        for i in range(1, 7):
            header_tags = soup.find_all(f'h{i}')
            for header in header_tags:
                text = header.get_text(strip=True)
                if text:
                    headers.append({
                        'level': i,
                        'text': text
                    })
        return headers

    def _extract_paragraphs(self, soup: BeautifulSoup) -> List[str]:
        """Extract paragraph content from the page."""
        paragraphs = []
        p_tags = soup.find_all('p')
        for p in p_tags:
            text = p.get_text(strip=True)
            if len(text) > 20 and not self._is_navigation_text(text):
                paragraphs.append(text)
        return paragraphs

    def _is_navigation_text(self, text: str) -> bool:
        """Check if text is likely navigation/menu text."""
        nav_indicators = ['home', 'about', 'contact', 'login', 'register', 'menu', 'search']
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in nav_indicators) and len(text) < 50

    def _extract_tables(self, soup: BeautifulSoup) -> List[Dict]:
        """Extract tables from the page."""
        tables = []
        table_tags = soup.find_all('table')
        for i, table in enumerate(table_tags):
            table_data = {
                'caption': '',
                'rows': []
            }
            caption = table.find('caption')
            table_data['caption'] = caption.get_text(strip=True) if caption else f'Table {i + 1}'
            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_data = [cell.get_text(strip=True) for cell in cells]
                if any(cell.strip() for cell in row_data):
                    table_data['rows'].append(row_data)
            if table_data['rows']:
                tables.append(table_data)
        return tables

    def _extract_links(self, soup: BeautifulSoup, base_url: str) -> List[str]:
        """Extract same-domain links from the page."""
        links = []
        link_tags = soup.find_all('a', href=True)
        for link in link_tags:
            href = link.get('href', '')
            if href:
                clean_url = self._clean_url(href, base_url)
                if clean_url and self._is_same_domain(clean_url):
                    if clean_url not in self.visited_urls and clean_url not in links:
                        links.append(clean_url)
        return links

    def crawl(self, start_url: str) -> List[Dict]:
        """Main crawling method to crawl all pages in the same domain."""
        try:
            self.base_domain = urlparse(start_url).netloc
            logger.info(f"Starting crawl for domain: {self.base_domain}")
            logger.info(f"Delay: {self.delay}s, Timeout: {self.timeout}s")

            self.pending_urls.append(start_url)
            while self.pending_urls and len(self.visited_urls) < self.max_pages:
                current_url = self.pending_urls.pop(0)
                if current_url in self.visited_urls:
                    continue

                logger.info(f"Crawling: {current_url} ({len(self.visited_urls) + 1}/{self.max_pages} pages)")
                page_content = self._extract_page_content(current_url)
                self.extracted_content.append(page_content)
                self.visited_urls.add(current_url)

                # Only add more URLs if we haven't reached the limit
                if len(self.visited_urls) < self.max_pages:
                    for link in page_content['links']:
                        if link not in self.visited_urls and link not in self.pending_urls:
                            self.pending_urls.append(link)

                if self.delay > 0:
                    time.sleep(self.delay)

            logger.info(f"Crawling completed. Visited {len(self.visited_urls)} pages.")
            return self.extracted_content
        except Exception as e:
            logger.error(f"Crawler error: {e}")
            return self.extracted_content

    def generate_word_document(self, output_path: str = 'crawled_content.docx') -> str:
        """Generate Word document from extracted content."""
        if not self.extracted_content:
            raise ValueError("No content extracted. Run crawl() first.")

        logger.info("Generating Word document...")
        doc = Document()
        title = doc.add_heading(f'Web Crawl Results - {self.base_domain}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary = doc.add_paragraph(f'Crawled {len(self.extracted_content)} pages from {self.base_domain}')
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for page_data in self.extracted_content:
            doc.add_page_break()
            doc.add_heading(page_data['url'], level=1)
            if page_data['title']:
                doc.add_heading(f"Title: {page_data['title']}", level=2)
            if page_data['headers']:
                doc.add_heading('Headers:', level=2)
                for header in page_data['headers']:
                    heading_level = min(header['level'] + 2, 9)
                    doc.add_heading(header['text'], level=heading_level)
            if page_data['paragraphs']:
                doc.add_heading('Content:', level=2)
                for paragraph in page_data['paragraphs']:
                    doc.add_paragraph(paragraph)
            if page_data['tables']:
                doc.add_heading('Tables:', level=2)
                for table_data in page_data['tables']:
                    doc.add_paragraph(table_data['caption'], style='Caption')
                    if table_data['rows']:
                        max_cols = max(len(row) for row in table_data['rows'])
                        table = doc.add_table(rows=len(table_data['rows']), cols=max_cols)
                        table.style = 'Table Grid'
                        for row_idx, row_data in enumerate(table_data['rows']):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < max_cols:
                                    table.cell(row_idx, col_idx).text = str(cell_data)
                    doc.add_paragraph()

        doc.save(output_path)
        logger.info(f"Word document saved to: {output_path}")
        return output_path

    def save_json(self, output_path: str = 'crawled_content.json') -> str:
        """Save extracted content as JSON file."""
        if not self.extracted_content:
            raise ValueError("No content extracted. Run crawl() first.")

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.extracted_content, f, indent=2, ensure_ascii=False)
        logger.info(f"JSON file saved to: {output_path}")
        return output_path

def generate_filename(base_name: str, extension: str, job_id: str) -> str:
    """Generate unique filename with timestamp and job ID."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base_name}_{timestamp}_{job_id[:8]}.{extension}"

async def crawl_website_background(
    start_url: str, 
    delay: float, 
    timeout: int, 
    max_pages: int,
    job_id: str,
    output_filename: Optional[str] = None
):
    """Background task to crawl website."""
    try:
        job_status[job_id] = {
            'status': 'running',
            'message': 'Crawling in progress...',
            'start_time': datetime.now()
        }

        crawler = WebCrawler(delay=delay, timeout=timeout, max_pages=1)
        logger.info(f"Job {job_id}: Starting crawl for {start_url}")
        content = crawler.crawl(start_url)

        if output_filename:
            base_name = output_filename.replace('.docx', '').replace('.json', '')
        else:
            domain = urlparse(start_url).netloc.replace('www.', '')
            base_name = f"crawl_{domain}"

        word_filename = generate_filename(base_name, 'docx', job_id)
        json_filename = generate_filename(base_name, 'json', job_id)

        word_path = crawler.generate_word_document(word_filename)
        json_path = crawler.save_json(json_filename)

        combined_content = " ".join(
            [h['text'] for doc in content for h in doc['headers']] +
            [p for doc in content for p in doc['paragraphs']]
        ) if content else ""

        job_status[job_id] = {
            'status': 'completed',
            'message': 'Crawling completed successfully',
            'pages_crawled': len(content),
            'word_document': word_path,
            'json_file': json_path,
            'end_time': datetime.now(),
            'content': combined_content
        }

        logger.info(f"Job {job_id}: Completed successfully. Crawled {len(content)} pages")
    except Exception as e:
        logger.error(f"Job {job_id}: Failed with error: {e}")
        job_status[job_id] = {
            'status': 'failed',
            'message': f'Crawling failed: {str(e)}',
            'end_time': datetime.now()
        }

@app.post("/crawl", response_model=CrawlResponse)
async def start_crawl(request: CrawlRequest, background_tasks: BackgroundTasks):
    """Start a web crawling job."""
    job_id = str(uuid.uuid4())
    background_tasks.add_task(
        crawl_website_background,
        str(request.start_url),
        request.delay,
        request.timeout,
        request.max_pages,
        job_id,
        request.output_filename
    )
    job_status[job_id] = {
        'status': 'started',
        'message': 'Crawling job started',
        'start_time': datetime.now()
    }
    return CrawlResponse(
        job_id=job_id,
        status='started',
        message='Crawling job started successfully'
    )

@app.get("/status/{job_id}", response_model=CrawlResponse)
async def get_crawl_status(job_id: str):
    """Get the status of a crawling job."""
    if job_id not in job_status:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job_status[job_id]
    return CrawlResponse(
        job_id=job_id,
        status=status['status'],
        message=status['message'],
        pages_crawled=status.get('pages_crawled'),
        word_document=status.get('word_document'),
        json_file=status.get('json_file')
    )

@app.get("/download/word/{job_id}")
async def download_word_document(job_id: str):
    """Download the Word document for a completed job."""
    if job_id not in job_status:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job_status[job_id]
    if status['status'] != 'completed':
        raise HTTPException(status_code=400, detail="Job not completed yet")
    word_path = status.get('word_document')
    if not word_path or not os.path.exists(word_path):
        raise HTTPException(status_code=404, detail="Word document not found")
    return FileResponse(
        path=word_path,
        filename=os.path.basename(word_path),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.get("/download/json/{job_id}")
async def download_json_file(job_id: str):
    """Download the JSON file for a completed job."""
    if job_id not in job_status:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job_status[job_id]
    if status['status'] != 'completed':
        raise HTTPException(status_code=400, detail="Job not completed yet")
    json_path = status.get('json_file')
    if not json_path or not os.path.exists(json_path):
        raise HTTPException(status_code=404, detail="JSON file not found")
    return FileResponse(
        path=json_path,
        filename=os.path.basename(json_path),
        media_type='application/json'
    )

@app.get("/jobs")
async def list_jobs():
    """List all crawling jobs."""
    return {
        "jobs": [
            {
                "job_id": job_id,
                "status": status["status"],
                "message": status["message"],
                "pages_crawled": status.get("pages_crawled"),
                "start_time": status.get("start_time").isoformat() if status.get("start_time") else None
            }
            for job_id, status in job_status.items()
        ]
    }

@app.delete("/jobs/{job_id}")
async def delete_job(job_id: str):
    """Delete a job and its associated files."""
    if job_id not in job_status:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job_status[job_id]
    for file_key in ['word_document', 'json_file']:
        file_path = status.get(file_key)
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Deleted file: {file_path}")
            except Exception as e:
                logger.error(f"Failed to delete file {file_path}: {e}")
    del job_status[job_id]
    return {"message": f"Job {job_id} deleted successfully"}

@app.get("/")
async def root():
    """API information."""
    return {
        "message": "Web Crawler API",
        "version": "1.0.0",
        "endpoints": {
            "POST /crawl": "Start a new crawling job",
            "GET /status/{job_id}": "Get job status",
            "GET /download/word/{job_id}": "Download Word document",
            "GET /download/json/{job_id}": "Download JSON file",
            "GET /jobs": "List all jobs",
            "DELETE /jobs/{job_id}": "Delete a job"
        }
    }

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)